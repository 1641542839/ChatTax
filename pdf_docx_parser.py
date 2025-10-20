"""
Parse PDF and DOCX files and extract text and tables.
Saves parsed content as JSON files.
"""

import os
import json
from pathlib import Path
from tqdm import tqdm
import PyPDF2
from docx import Document
import pdfplumber


class DocumentParser:
    """Parser for PDF and DOCX documents."""
    
    def __init__(self, input_dir="./data/raw_files", output_dir="./data/parsed", table_dir="./data/tables"):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.table_dir = table_dir
        
    def parse_pdf(self, pdf_path):
        """
        Parse PDF file and extract text and tables.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Dictionary with parsed content
        """
        content = {
            'text': '',
            'tables': [],
            'num_pages': 0
        }
        
        try:
            # Extract text using PyPDF2
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                content['num_pages'] = len(pdf_reader.pages)
                
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                
                content['text'] = '\n'.join(text_parts)
            
            # Extract tables using pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            content['tables'].append({
                                'page': page_num + 1,
                                'table_num': table_num + 1,
                                'data': table
                            })
        
        except Exception as e:
            print(f"Error parsing PDF {pdf_path}: {str(e)}")
        
        return content
    
    def parse_docx(self, docx_path):
        """
        Parse DOCX file and extract text and tables.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary with parsed content
        """
        content = {
            'text': '',
            'tables': [],
            'num_paragraphs': 0
        }
        
        try:
            doc = Document(docx_path)
            
            # Extract text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            content['text'] = '\n'.join(paragraphs)
            content['num_paragraphs'] = len(paragraphs)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                content['tables'].append({
                    'table_num': table_num + 1,
                    'data': table_data
                })
        
        except Exception as e:
            print(f"Error parsing DOCX {docx_path}: {str(e)}")
        
        return content
    
    def save_tables_as_csv(self, tables, source_filename):
        """
        Save extracted tables as CSV files.
        
        Args:
            tables: List of table dictionaries
            source_filename: Original filename (without extension)
        """
        os.makedirs(self.table_dir, exist_ok=True)
        
        for table_info in tables:
            table_data = table_info['data']
            if not table_data:
                continue
            
            # Create filename for this table
            if 'page' in table_info:
                csv_filename = f"{source_filename}_page{table_info['page']}_table{table_info['table_num']}.csv"
            else:
                csv_filename = f"{source_filename}_table{table_info['table_num']}.csv"
            
            csv_path = os.path.join(self.table_dir, csv_filename)
            
            # Write CSV
            try:
                with open(csv_path, 'w', encoding='utf-8') as f:
                    for row in table_data:
                        # Escape commas and quotes in cells
                        escaped_row = [f'"{cell}"' if ',' in cell or '"' in cell else cell 
                                     for cell in row]
                        f.write(','.join(escaped_row) + '\n')
            except Exception as e:
                print(f"Error saving table {csv_filename}: {str(e)}")
    
    def process_file(self, file_path):
        """
        Process a single PDF or DOCX file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            content = self.parse_pdf(file_path)
        elif file_ext == '.docx':
            content = self.parse_docx(file_path)
        else:
            print(f"Unsupported file type: {file_ext}")
            return None
        
        # Create metadata
        metadata = {
            'source_file': str(file_path),
            'file_type': file_ext[1:],  # Remove dot
            'filename': Path(file_path).name
        }
        
        # Save tables as CSV
        if content.get('tables'):
            self.save_tables_as_csv(content['tables'], Path(file_path).stem)
        
        return {
            'metadata': metadata,
            'content': content['text'],
            'tables_count': len(content.get('tables', [])),
            'content_type': 'text'
        }
    
    def process_directory(self):
        """Process all PDF and DOCX files in the input directory."""
        os.makedirs(self.output_dir, exist_ok=True)
        
        if not os.path.exists(self.input_dir):
            print(f"Input directory does not exist: {self.input_dir}")
            return
        
        # Find all PDF and DOCX files
        file_patterns = ['*.pdf', '*.docx']
        files = []
        for pattern in file_patterns:
            files.extend(Path(self.input_dir).glob(pattern))
        
        if not files:
            print("No PDF or DOCX files found")
            return
        
        print(f"Processing {len(files)} documents...")
        
        for file_path in tqdm(files, desc="Parsing documents"):
            try:
                parsed_data = self.process_file(file_path)
                
                if parsed_data:
                    # Save as JSON
                    output_filename = file_path.stem + '.json'
                    output_path = os.path.join(self.output_dir, output_filename)
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(parsed_data, f, indent=2, ensure_ascii=False)
                        
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")


def main():
    """Main function to parse documents."""
    parser = DocumentParser()
    
    print("Starting document parsing...")
    parser.process_directory()
    print("Parsing completed!")


if __name__ == "__main__":
    main()
